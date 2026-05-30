import os
import json
import shutil

from fastapi import FastAPI, File, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document
from langchain_core.vectorstores import InMemoryVectorStore
from langchain_ollama import OllamaEmbeddings
import ollama

load_dotenv()

OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "mistral")
OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
OLLAMA_EMBED_MODEL = os.getenv("OLLAMA_EMBED_MODEL", "nomic-embed-text")
ollama_client = ollama.Client(host=OLLAMA_BASE_URL)

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

vectorstore = None
document_text = ""


class Query(BaseModel):
    question: str
    language: str = "hi"


class TemplateRequest(BaseModel):
    template_type: str
    fields: dict
    language: str = "hi"


class ExplainRequest(BaseModel):
    text: str
    language: str = "hi"


class CustomDocRequest(BaseModel):
    description: str
    parties: str = ""
    key_terms: str = ""
    jurisdiction: str = "India"
    language: str = "hi"


class CompareRequest(BaseModel):
    doc_a_name: str
    doc_a_analysis: dict
    doc_b_name: str
    doc_b_analysis: dict
    language: str = "hi"


def get_message_content(response) -> str:
    if isinstance(response, dict):
        return response["message"]["content"].strip()
    message = getattr(response, "message", None)
    if isinstance(message, dict):
        return message.get("content", "").strip()
    return getattr(message, "content", "").strip()


def get_model_name(model) -> str:
    if isinstance(model, dict):
        return model.get("model") or model.get("name") or ""
    return getattr(model, "model", None) or getattr(model, "name", None) or ""


def call_ollama_json(prompt: str) -> dict:
    try:
        response = ollama_client.chat(
            model=OLLAMA_MODEL,
            messages=[{"role": "user", "content": prompt}],
            format="json",
            options={"temperature": 0},
        )
        return json.loads(get_message_content(response))
    except Exception as e:
        print(f"Ollama JSON call failed: {e}")
        return {}


def call_ollama_text(prompt: str) -> str:
    try:
        response = ollama_client.chat(
            model=OLLAMA_MODEL,
            messages=[{"role": "user", "content": prompt}],
            options={"temperature": 0},
        )
        return get_message_content(response)
    except Exception as e:
        print(f"Ollama text call failed: {e}")
        return "उत्तर देने में त्रुटि हुई। कृपया पुनः प्रयास करें।"


def load_document(file_path: str) -> list:
    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext == ".pdf":
            return PyPDFLoader(file_path).load()
        elif ext in (".txt", ".md"):
            return TextLoader(file_path, encoding="utf-8").load()
        elif ext == ".docx":
            import docx
            doc = docx.Document(file_path)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return [Document(page_content=text, metadata={"source": file_path})]
        else:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return [Document(page_content=f.read(), metadata={"source": file_path})]
    except Exception as e:
        print(f"Document load error: {e}")
        return [Document(page_content="", metadata={"source": file_path})]


async def analyze_document(text: str) -> dict:
    prompt = f"""You are an expert legal document analyzer. Analyze the legal document below and return ONLY a valid JSON object — no extra text, no markdown.

Document:
{text[:3500]}

Return exactly this JSON structure:
{{
  "document_type": "Type of legal document (e.g. Rental Agreement, NDA, Service Contract, Employment Agreement, Power of Attorney)",
  "parties": ["Party 1 name", "Party 2 name"],
  "key_dates": ["Important date 1 with context", "Important date 2 with context"],
  "key_clauses": ["Most important obligation or clause 1", "Clause 2", "Clause 3"],
  "risk_score": 45,
  "risk_level": "Medium",
  "risk_factors": ["Specific risky clause or term 1", "Risk 2"],
  "suggested_questions": [
    "Question 1 relevant to this document?",
    "Question 2?",
    "Question 3?",
    "Question 4?",
    "Question 5?"
  ]
}}

Rules:
- risk_score is an integer 0-100 (0-30=Low, 31-60=Medium, 61-100=High)
- risk_level must be exactly "Low", "Medium", or "High"
- suggested_questions must be specific to this document type
- Return ONLY the JSON object"""

    result = call_ollama_json(prompt)

    if not result:
        return {
            "document_type": "Legal Document",
            "parties": ["Party not identified"],
            "key_dates": ["No specific dates found"],
            "key_clauses": ["Document processed — ask questions to explore clauses"],
            "risk_score": 40,
            "risk_level": "Medium",
            "risk_factors": ["Manual review recommended"],
            "suggested_questions": [
                "मुख्य जोखिम क्या हैं?",
                "समय सीमा कब तक है?",
                "भुगतान की शर्तें क्या हैं?",
                "अनुबंध तोड़ने पर क्या होगा?",
                "दस्तावेज़ की मुख्य शर्तें क्या हैं?",
            ],
        }

    result.setdefault("document_type", "Legal Document")
    result.setdefault("parties", [])
    result.setdefault("key_dates", [])
    result.setdefault("key_clauses", [])
    result.setdefault("risk_score", 40)
    result.setdefault("risk_level", "Medium")
    result.setdefault("risk_factors", [])
    result.setdefault("suggested_questions", [])
    return result


async def verify_document_text(text: str) -> dict:
    prompt = f"""You are a legal document verification expert. Analyze this document for issues and return ONLY a valid JSON object.

Document:
{text[:3500]}

Return exactly this JSON structure:
{{
  "missing_signatures": ["Specific finding about missing or incomplete signatures, or empty list if none"],
  "inconsistent_dates": ["Specific date inconsistency found with details, or empty list if none"],
  "clause_mismatches": ["Specific problematic or unusual clause with details"],
  "overall_status": "Valid",
  "recommendations": ["Actionable recommendation 1", "Recommendation 2", "Recommendation 3"]
}}

overall_status must be one of: "Valid", "Issues Found", "Needs Review"
Return ONLY the JSON object."""

    result = call_ollama_json(prompt)

    if not result:
        return {
            "missing_signatures": [],
            "inconsistent_dates": [],
            "clause_mismatches": [],
            "overall_status": "Needs Review",
            "recommendations": ["दस्तावेज़ की मैन्युअल समीक्षा करें।"],
        }

    result.setdefault("missing_signatures", [])
    result.setdefault("inconsistent_dates", [])
    result.setdefault("clause_mismatches", [])
    result.setdefault("overall_status", "Needs Review")
    result.setdefault("recommendations", [])
    return result


@app.post("/upload/")
async def upload_file(file: UploadFile = File(...)):
    global vectorstore, document_text

    file_location = f"temp_{file.filename}"
    with open(file_location, "wb+") as file_object:
        shutil.copyfileobj(file.file, file_object)

    documents = load_document(file_location)

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    texts = text_splitter.split_documents(documents)
    document_text = "\n".join([doc.page_content for doc in texts[:20]])

    embeddings = OllamaEmbeddings(model=OLLAMA_EMBED_MODEL, base_url=OLLAMA_BASE_URL)
    vectorstore = InMemoryVectorStore.from_documents(texts, embeddings)

    os.remove(file_location)

    analysis = await analyze_document(document_text)
    verification = await verify_document_text(document_text)

    return {
        "message": "File uploaded and processed successfully",
        "filename": file.filename,
        "analysis": analysis,
        "verification": verification,
    }


@app.post("/ask/")
async def ask_question(query: Query):
    global vectorstore
    if not vectorstore:
        return {"error": "Please upload a document first"}

    retriever = vectorstore.as_retriever(search_kwargs={"k": 4})
    docs = retriever.invoke(query.question)
    context = "\n\n".join([doc.page_content for doc in docs])

    lang_instruction = (
        "Answer in simple Hindi. Use Devanagari script."
        if query.language == "hi"
        else "Answer in simple English."
    )
    prompt = f"""You are a helpful legal assistant. Use the context below to answer the question.
{lang_instruction}
If you don't know the answer from the context, say so — don't make things up.
Use **bold** for important terms and - for bullet points where appropriate.

Context:
{context}

Question: {query.question}

Answer:"""

    answer = call_ollama_text(prompt)
    return {
        "answer": answer,
        "sources": [
            {"content": doc.page_content, "page": doc.metadata.get("page_number", 1)}
            for doc in docs
        ],
    }


@app.post("/explain/")
async def explain_clause(req: ExplainRequest):
    lang_instruction = (
        "Explain in simple Hindi (Devanagari script) in 2-3 sentences that a non-lawyer can understand."
        if req.language == "hi"
        else "Explain in simple English in 2-3 sentences that a non-lawyer can understand."
    )
    prompt = f"""You are a legal expert. {lang_instruction}

Clause or term to explain:
"{req.text}"

Plain-language explanation:"""
    return {"explanation": call_ollama_text(prompt)}


@app.post("/generate/")
async def generate_template(req: TemplateRequest):
    lang_instruction = (
        "in simple Hindi using Devanagari script. Make it formal and legally appropriate for India."
        if req.language == "hi"
        else "in simple English. Make it formal and legally appropriate for India."
    )
    fields_text = "\n".join(f"- {k}: {v}" for k, v in req.fields.items() if v)
    prompt = f"""You are an expert Indian legal document drafter. Generate a complete {req.template_type} {lang_instruction}.

Use the following details:
{fields_text}

Formatting rules (strictly follow):
- Use # for the document title (e.g., # RENTAL AGREEMENT)
- Use ## for major section headings (e.g., ## 1. PARTIES, ## 2. TERMS)
- Use ### for sub-sections
- Use numbered lists (1. 2. 3.) for clauses
- Use **bold** only for party names, dates, and key amounts
- End with a ## SIGNATURES section with blank lines for signing
- Do NOT write "Here is the document" or any preamble — start directly with the title

Generate the complete legal document:"""

    return {"document": call_ollama_text(prompt)}


@app.post("/generate-custom/")
async def generate_custom_document(req: CustomDocRequest):
    lang_instruction = (
        "in Hindi using Devanagari script. Make it formal and legally appropriate for India."
        if req.language == "hi"
        else "in English. Make it formal and legally appropriate for India."
    )
    parts = [f"Document description: {req.description}"]
    if req.parties.strip():
        parts.append(f"Parties involved: {req.parties}")
    if req.key_terms.strip():
        parts.append(f"Key terms and requirements: {req.key_terms}")
    if req.jurisdiction.strip():
        parts.append(f"Jurisdiction: {req.jurisdiction}")
    details = "\n".join(f"- {p}" for p in parts)

    prompt = f"""You are an expert Indian legal document drafter. Based on the user's description below, generate a complete, professional legal document {lang_instruction}.

User requirements:
{details}

Formatting rules (strictly follow):
- Use # for the document title in ALL CAPS (e.g., # LOAN AGREEMENT)
- Use ## for major sections (e.g., ## 1. PARTIES INVOLVED, ## 2. TERMS AND CONDITIONS)
- Use ### for sub-sections
- Use numbered lists for clauses within sections
- Use **bold** only for names, key dates, and monetary amounts
- Include ## RECITALS, ## DEFINITIONS, ## OBLIGATIONS, ## TERMINATION, ## DISPUTE RESOLUTION, ## GOVERNING LAW, ## SIGNATURES sections
- End with a ## SIGNATURES section with blank signature lines
- Do NOT write any preamble — start directly with the # title

Generate the complete legal document now:"""

    result = call_ollama_text(prompt)
    return {"document": result, "detected_type": req.description[:80]}


@app.post("/ai-compare/")
async def ai_compare(req: CompareRequest):
    lang_instruction = "Respond in simple Hindi using Devanagari script." if req.language == "hi" else "Respond in simple English."
    a = req.doc_a_analysis
    b = req.doc_b_analysis

    prompt = f"""You are a legal document expert. Compare these two documents and give a clear recommendation.
{lang_instruction}

Document A: {req.doc_a_name}
- Type: {a.get('document_type', 'Unknown')}
- Risk Score: {a.get('risk_score', 0)}/100 ({a.get('risk_level', 'Unknown')})
- Risk Factors: {', '.join(a.get('risk_factors', [])[:3])}
- Key Clauses: {', '.join(a.get('key_clauses', [])[:3])}

Document B: {req.doc_b_name}
- Type: {b.get('document_type', 'Unknown')}
- Risk Score: {b.get('risk_score', 0)}/100 ({b.get('risk_level', 'Unknown')})
- Risk Factors: {', '.join(b.get('risk_factors', [])[:3])}
- Key Clauses: {', '.join(b.get('key_clauses', [])[:3])}

Provide:
1. Which document is better and why (2-3 specific reasons)
2. Key risks in each document that the user should watch out for
3. A final recommendation in one sentence

Be specific, practical, and use simple language a non-lawyer can understand."""

    result = call_ollama_text(prompt)
    return {"feedback": result}


@app.get("/verify/")
async def verify_document():
    global vectorstore, document_text

    if not vectorstore or not document_text:
        return {
            "report": {
                "missing_signatures": ["कोई दस्तावेज़ अपलोड नहीं है।"],
                "inconsistent_dates": [],
                "clause_mismatches": [],
                "overall_status": "No Document",
                "recommendations": ["दस्तावेज़ अपलोड करके पुनः प्रयास करें।"],
            }
        }

    return {"report": await verify_document_text(document_text)}


@app.get("/")
async def root():
    return {
        "status": "Nyay-Saarthi backend running",
        "llm_provider": "ollama",
        "ollama_base_url": OLLAMA_BASE_URL,
        "model": OLLAMA_MODEL,
        "embedding_model": OLLAMA_EMBED_MODEL,
    }


@app.get("/health/")
async def health():
    try:
        response = ollama_client.list()
        models = response.get("models", []) if isinstance(response, dict) else getattr(response, "models", [])
        return {
            "status": "ok",
            "llm_provider": "ollama",
            "ollama_base_url": OLLAMA_BASE_URL,
            "model": OLLAMA_MODEL,
            "embedding_model": OLLAMA_EMBED_MODEL,
            "available_models": [name for name in (get_model_name(m) for m in models) if name],
        }
    except Exception as e:
        return {
            "status": "ollama_unavailable",
            "llm_provider": "ollama",
            "ollama_base_url": OLLAMA_BASE_URL,
            "model": OLLAMA_MODEL,
            "embedding_model": OLLAMA_EMBED_MODEL,
            "error": str(e),
        }


if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", 8000))
    host = os.getenv("HOST", "0.0.0.0")
    uvicorn.run(app, host=host, port=port)
